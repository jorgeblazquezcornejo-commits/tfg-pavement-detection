import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =======================
# 1) Datos
# =======================
data = [
    ["China_Drone", "China_Drone", 0.636197, 0.377576, 0.678294, 0.610938, 0.642857],
    ["China_Drone", "China_MotorBike", 0.960559, 0.684083, 0.922749, 0.918662, 0.920701],
    ["China_Drone", "Czech", 0.676505, 0.357321, 0.688871, 0.616351, 0.650596],
    ["China_Drone", "India", 0.680758, 0.387249, 0.695551, 0.614348, 0.652433],
    ["China_Drone", "Japan", 0.79705, 0.467497, 0.769672, 0.736381, 0.752659],
    ["China_Drone", "Norway", 0.41867, 0.225048, 0.598799, 0.394915, 0.475941],
    ["China_Drone", "United_States", 0.781583, 0.523329, 0.743317, 0.745444, 0.744379],

    ["China_MotorBike", "China_Drone", 0.853263, 0.576218, 0.831909, 0.806749, 0.819136],
    ["China_MotorBike", "China_MotorBike", 0.37509, 0.17741, 0.482915, 0.432775, 0.456472],
    ["China_MotorBike", "Czech", 0.691519, 0.368752, 0.682803, 0.657205, 0.669759],
    ["China_MotorBike", "India", 0.672977, 0.381866, 0.727361, 0.610706, 0.663948],
    ["China_MotorBike", "Japan", 0.791309, 0.460795, 0.7691, 0.729044, 0.748536],
    ["China_MotorBike", "Norway", 0.431234, 0.231499, 0.613158, 0.394925, 0.480419],
    ["China_MotorBike", "United_States", 0.788865, 0.52886, 0.781506, 0.708808, 0.743384],

    ["Czech", "China_Drone", 0.88971, 0.624906, 0.87005, 0.827208, 0.848088],
    ["Czech", "China_MotorBike", 0.961789, 0.697612, 0.925612, 0.936769, 0.931157],
    ["Czech", "Czech", 0.168414, 0.059743, 0.27256, 0.254208, 0.263064],
    ["Czech", "India", 0.721626, 0.431299, 0.730037, 0.64785, 0.686492],
    ["Czech", "Japan", 0.830324, 0.50399, 0.810789, 0.757122, 0.783037],
    ["Czech", "Norway", 0.45679, 0.255502, 0.63455, 0.419788, 0.505296],
    ["Czech", "United_States", 0.814468, 0.562735, 0.790183, 0.752655, 0.770963],

    ["India", "China_Drone", 0.88276, 0.613002, 0.831385, 0.833288, 0.832335],
    ["India", "China_MotorBike", 0.961014, 0.690411, 0.926018, 0.932228, 0.929113],
    ["India", "Czech", 0.695736, 0.388592, 0.702709, 0.65232, 0.676578],
    ["India", "India", 0.099334, 0.040773, 0.273233, 0.152217, 0.195514],
    ["India", "Japan", 0.815695, 0.491003, 0.79085, 0.748713, 0.769205],
    ["India", "Norway", 0.439942, 0.245174, 0.621129, 0.401912, 0.488033],
    ["India", "United_States", 0.797659, 0.543406, 0.777719, 0.736812, 0.756713],

    ["Japan", "China_Drone", 0.899226, 0.639032, 0.873869, 0.840779, 0.857005],
    ["Japan", "China_MotorBike", 0.960487, 0.692381, 0.928071, 0.925681, 0.926875],
    ["Japan", "Czech", 0.705642, 0.401616, 0.736297, 0.621265, 0.673908],
    ["Japan", "India", 0.727314, 0.438633, 0.73477, 0.666637, 0.699047],
    ["Japan", "Japan", 0.331431, 0.148128, 0.435862, 0.366761, 0.398337],
    ["Japan", "Norway", 0.459413, 0.257464, 0.652225, 0.417925, 0.509426],
    ["Japan", "United_States", 0.798051, 0.543476, 0.780434, 0.734744, 0.7569],

    ["Norway", "China_Drone", 0.894918, 0.629542, 0.870321, 0.832744, 0.851118],
    ["Norway", "China_MotorBike", 0.968603, 0.695949, 0.930016, 0.930361, 0.930188],
    ["Norway", "Czech", 0.687004, 0.375278, 0.711866, 0.61834, 0.661815],
    ["Norway", "India", 0.717274, 0.418373, 0.725606, 0.658865, 0.690627],
    ["Norway", "Japan", 0.813045, 0.48283, 0.77843, 0.748468, 0.763155],
    ["Norway", "Norway", 0.118914, 0.046228, 0.315093, 0.153413, 0.206356],
    ["Norway", "United_States", 0.810464, 0.549876, 0.79633, 0.746039, 0.770365],

    ["United_States", "China_Drone", 0.879146, 0.614091, 0.847263, 0.832514, 0.839824],
    ["United_States", "China_MotorBike", 0.964592, 0.693236, 0.924592, 0.931902, 0.928232],
    ["United_States", "Czech", 0.710511, 0.397438, 0.730997, 0.65539, 0.691132],
    ["United_States", "India", 0.704036, 0.414736, 0.712547, 0.645965, 0.677624],
    ["United_States", "Japan", 0.820149, 0.493889, 0.800666, 0.749658, 0.774323],
    ["United_States", "Norway", 0.445282, 0.2461, 0.619402, 0.411745, 0.494664],
    ["United_States", "United_States", 0.461617, 0.251995, 0.506813, 0.502044, 0.504417],
]

internal_cols = ["Modelo", "Dataset_test", "mAP50", "mAP50-95", "Precision_media", "Recall_media", "F1_media"]
df = pd.DataFrame(data, columns=internal_cols)

display_cols = [
    "Dataset no incluido en el entrenamiento del modelo",
    "Dataset utilizado en el test",
    "mAP50",
    "mAP50-95",
    "Precision media",
    "Recall media",
    "F1 media",
]

model_colors_hex = {
    "China_Drone": "#FFD1DC",
    "China_MotorBike": "#FFDEAD",
    "Czech": "#98FB98",
    "India": "#FFFACD",
    "Japan": "#E6E6FA",
    "Norway": "#ADD8E6",
    "United_States": "#D3D3D3",
}


# =======================
# 2) ORDENAR (lo que pediste)
#   - Excluir self-test (Dataset_test == Modelo)
#   - Para cada modelo: 6 filas ordenadas por F1_media desc
#   - Mantener un orden de modelos fijo (opcional pero recomendado)
# =======================
df = df[df["Dataset_test"] != df["Modelo"]].copy()

order_models = ["China_Drone", "China_MotorBike", "Czech", "India", "Japan", "Norway", "United_States"]
df["Modelo"] = pd.Categorical(df["Modelo"], categories=order_models, ordered=True)

df.sort_values(by=["Modelo", "F1_media"], ascending=[True, False], inplace=True)
df.reset_index(drop=True, inplace=True)


# =======================
# 3) Utilidades Word (colores + header repetido)
# =======================
def hex_to_word_fill(hex_color: str) -> str:
    """Convierte '#RRGGBB' a 'RRGGBB' para shading en Word."""
    return hex_color.strip().lstrip("#").upper()

def set_cell_shading(cell, fill: str):
    """Aplica color de fondo a una celda (Word). fill = 'RRGGBB'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_repeat_table_header(row):
    """Hace que el encabezado se repita en cada página."""
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


# =======================
# 4) Exportar a Word (TFG)
# =======================
doc = Document()

# A4 horizontal + márgenes estrechos
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# doc.add_heading("Resultados de validación Leave-One-Out", level=2)

table = doc.add_table(rows=1, cols=len(display_cols))
table.style = "Table Grid"

# Header
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(display_cols):
    hdr_cells[i].text = col_name
    for p in hdr_cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(8)
    set_cell_shading(hdr_cells[i], "404040")  # gris oscuro

set_repeat_table_header(table.rows[0])

# Filas
for _, r in df.iterrows():
    row_cells = table.add_row().cells

    values = [
        r["Modelo"],
        r["Dataset_test"],
        f"{r['mAP50']:.6f}",
        f"{r['mAP50-95']:.6f}",
        f"{r['Precision_media']:.6f}",
        f"{r['Recall_media']:.6f}",
        f"{r['F1_media']:.6f}",
    ]

    fill = hex_to_word_fill(model_colors_hex.get(str(r["Modelo"]), "#FFFFFF"))

    for i, val in enumerate(values):
        row_cells[i].text = str(val)
        set_cell_shading(row_cells[i], fill)
        for p in row_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8)

# Guardar Word
doc.save("tabla_resultados_TFG.docx")
print("OK: generado 'tabla_resultados_TFG.docx' (tabla editable, con saltos de página).")


# =======================
# 5) (Opcional) Generar PNG (figura)
# =======================
fig, ax = plt.subplots(figsize=(30, 10))
ax.axis("off")

display_cols_fig = [
    "Dataset no incluido en el\nentrenamiento del modelo",
    "Dataset utilizado en el\ntest",
    "mAP50",
    "mAP50-95",
    "Precision media",
    "Recall media",
    "F1 media",
]

col_widths = [0.40, 0.28, 0.09, 0.10, 0.12, 0.11, 0.10]

# IMPORTANTE: para la figura mostramos SOLO las columnas (no los headers internos)
cell_text = df[internal_cols].values

table_fig = ax.table(
    cellText=cell_text,
    colLabels=display_cols_fig,
    colWidths=col_widths,
    loc="center",
    cellLoc="center",
)

table_fig.auto_set_font_size(False)
table_fig.set_fontsize(9)

for (row, col), cell in table_fig.get_celld().items():
    if row == 0:
        cell.set_facecolor("#404040")
        cell.set_text_props(color="white", weight="bold")
        cell.get_text().set_wrap(True)
        cell.set_height(cell.get_height() * 1.8)
    else:
        model_name = str(df.iloc[row - 1, 0])
        cell.set_facecolor(model_colors_hex.get(model_name, "#FFFFFF"))

plt.savefig("tabla_resultados_TFG.png", dpi=300, bbox_inches="tight")
plt.show()

print("OK: generado 'tabla_resultados_TFG.png'.")
